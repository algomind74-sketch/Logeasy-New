from app.database.crud import get_logs
from app.database.db_session import SessionLocal
from datetime import datetime
import os
from docx import Document

COMPLIANCE_DIR = "data/reports/compliance/"
os.makedirs(COMPLIANCE_DIR, exist_ok=True)


def generate_compliance_report(start_date=None, end_date=None, filename=None):
    """
    Generate a compliance report (DOCX) summarizing logs within a date range.

    Args:
        start_date (datetime, optional): Filter logs from this date
        end_date (datetime, optional): Filter logs until this date
        filename (str, optional): Filename for DOCX report

    Returns:
        str: Path to the saved DOCX report
    """
    db = SessionLocal()
    logs = get_logs(db, start_date=start_date, end_date=end_date)
    db.close()

    if not filename:
        today_str = datetime.now().strftime("%Y_%m_%d")
        filename = f"compliance_report_{today_str}.docx"

    path = os.path.join(COMPLIANCE_DIR, filename)

    # Create DOCX report
    doc = Document()
    doc.add_heading("LogEasy Compliance Report", level=0)
    if start_date or end_date:
        doc.add_paragraph(
            f"Date Range: {start_date if start_date else 'Beginning'} - {end_date if end_date else 'Now'}"
        )
    doc.add_paragraph(f"Total Logs: {len(logs)}\n")

    # Summarize logs by level
    summary = {}
    for log in logs:
        summary[log.level] = summary.get(log.level, 0) + 1

    doc.add_heading("Log Summary by Level", level=1)
    for level, count in summary.items():
        doc.add_paragraph(f"{level}: {count}")

    # Optionally, list all logs
    doc.add_heading("Detailed Logs", level=1)
    for log in logs:
        doc.add_paragraph(f"[{log.created_at}] [{log.level}] {log.message}")

    doc.save(path)
    print(f"✅ Compliance report generated: {path}")
    return path


# Optional: test
if __name__ == "__main__":
    report_path = generate_compliance_report()
    print("Report saved at:", report_path)
