import os
from app.database.db_session import SessionLocal
from app.database.crud import get_logs
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

REPORT_DIR = "data/reports/"
os.makedirs(REPORT_DIR, exist_ok=True)

def generate_pdf_report(filename="report.pdf"):
    db = SessionLocal()
    logs = get_logs(db)
    db.close()
    
    path = os.path.join(REPORT_DIR, filename)
    c = canvas.Canvas(path, pagesize=letter)
    c.setFont("Helvetica", 12)
    y = 750
    c.drawString(50, y + 20, "LogEasy - Log Report")
    for log in logs:
        y -= 20
        c.drawString(50, y, f"[{log.level}] {log.message} - {log.created_at}")
        if y < 50:
            c.showPage()
            y = 750
    c.save()
    print(f"PDF report saved: {path}")

def generate_docx_report(filename="report.docx"):
    db = SessionLocal()
    logs = get_logs(db)
    db.close()
    
    doc = Document()
    doc.add_heading("LogEasy - Log Report", level=1)
    for log in logs:
        doc.add_paragraph(f"[{log.level}] {log.message} - {log.created_at}")
    path = os.path.join(REPORT_DIR, filename)
    doc.save(path)
    print(f"DOCX report saved: {path}")

if __name__ == "__main__":
    generate_pdf_report("report_2025_10_28.pdf")
    generate_docx_report("report_2025_10_28.docx")
